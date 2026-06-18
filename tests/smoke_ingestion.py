"""Smoke test for Step 2 ingestion: generate a .docx, .pdf and .txt with the same
clauses, run load_document on each, and print what came back.

Run:  python tests/smoke_ingestion.py
"""

from __future__ import annotations

import sys
from pathlib import Path

ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT / "src"))

from legalcompare.ingestion.loader import load_document  # noqa: E402
from legalcompare.ingestion.ocr import tesseract_available  # noqa: E402

OUT = ROOT / "data" / "samples" / "_ingest_demo"
OUT.mkdir(parents=True, exist_ok=True)

CLAUSES = [
    "1. Confidentiality. Each party shall keep the other party's Confidential "
    "Information strictly confidential and shall not disclose it to any third party.",
    "2. Term. This Agreement commences on the Effective Date and continues for "
    "twelve (12) months unless terminated earlier in accordance with its terms.",
    "3. Governing Law. This Agreement shall be governed by the laws of Singapore.",
]


def make_docx(path: Path) -> None:
    import docx

    d = docx.Document()
    d.add_heading("Mutual Non-Disclosure Agreement", level=1)
    for c in CLAUSES:
        d.add_paragraph(c)
    d.save(str(path))


def make_pdf(path: Path) -> None:
    from reportlab.lib.pagesizes import LETTER
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

    doc = SimpleDocTemplate(str(path), pagesize=LETTER)
    styles = getSampleStyleSheet()
    flow = [Paragraph("Mutual Non-Disclosure Agreement", styles["Title"]), Spacer(1, 12)]
    for c in CLAUSES:
        flow += [Paragraph(c, styles["BodyText"]), Spacer(1, 8)]
    doc.build(flow)


def make_txt(path: Path) -> None:
    path.write_text("\n".join(CLAUSES), encoding="utf-8")


def report(label: str, path: Path) -> None:
    doc = load_document(path)
    print(f"\n[{label}]  {path.name}")
    print(f"  blocks={len(doc.blocks)}  used_ocr={doc.used_ocr}  chars={len(doc.full_text)}")
    first = doc.blocks[0] if doc.blocks else None
    if first:
        preview = first.text[:80].replace("\n", " ")
        print(f"  block[0] page={first.page} offset=({first.start_char},{first.end_char})")
        print(f"  block[0] text: {preview}...")


def main() -> None:
    docx_path, pdf_path, txt_path = OUT / "nda.docx", OUT / "nda.pdf", OUT / "nda.txt"
    make_docx(docx_path)
    make_pdf(pdf_path)
    make_txt(txt_path)

    print(f"Tesseract available for OCR: {tesseract_available()}")
    report("DOCX", docx_path)
    report("PDF (text-based)", pdf_path)
    report("TXT", txt_path)
    print("\nSmoke test OK.")


if __name__ == "__main__":
    main()
