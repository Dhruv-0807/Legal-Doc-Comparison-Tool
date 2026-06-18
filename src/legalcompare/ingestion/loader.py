"""Step 2: read a file into a `Document` (full text + provenance-tagged blocks).

Dispatch by extension:
    .pdf   -> pdfplumber for text; per page, fall back to OCR if the page has
              (almost) no extractable text (i.e. it's scanned/image-only).
    .docx  -> python-docx (already digital text; no OCR needed).
    .txt   -> read as-is.

Why provenance matters: every block records its page and character offsets into
`full_text`, and whether it came from OCR. Later stages carry those offsets so any
flagged change can be traced back to the exact source span — a hard requirement.
"""

from __future__ import annotations

import logging
from pathlib import Path

from ..config import get_config
from ..schemas import Document, SourceBlock
from .ocr import ocr_image, tesseract_available

logger = logging.getLogger(__name__)


def load_document(path: str | Path) -> Document:
    """Read a supported file into a Document."""
    p = Path(path)
    if not p.exists():
        raise FileNotFoundError(p)

    ext = p.suffix.lower()
    if ext == ".pdf":
        return _load_pdf(p)
    if ext == ".docx":
        return _load_docx(p)
    if ext in {".txt", ".md"}:
        return _load_text(p)
    raise ValueError(f"Unsupported file type: {ext!r} (supported: .pdf, .docx, .txt)")


# --------------------------------------------------------------------------- #
# Helpers: assemble full_text + offsets consistently across all readers.
# --------------------------------------------------------------------------- #
def _assemble(name: str, raw_blocks: list[dict]) -> Document:
    """Join block texts into full_text and record exact char offsets per block.

    `raw_blocks` items: {"text": str, "page": int|None, "from_ocr": bool}.
    Blocks are joined with a single newline; offsets index into the joined text.
    """
    blocks: list[SourceBlock] = []
    parts: list[str] = []
    cursor = 0
    used_ocr = False

    for rb in raw_blocks:
        text = rb["text"].strip()
        if not text:
            continue
        start = cursor
        end = start + len(text)
        blocks.append(
            SourceBlock(
                text=text,
                page=rb.get("page"),
                start_char=start,
                end_char=end,
                from_ocr=rb.get("from_ocr", False),
            )
        )
        parts.append(text)
        cursor = end + 1  # +1 for the "\n" joiner
        used_ocr = used_ocr or rb.get("from_ocr", False)

    return Document(
        name=name,
        full_text="\n".join(parts),
        blocks=blocks,
        used_ocr=used_ocr,
    )


def _load_text(p: Path) -> Document:
    text = p.read_text(encoding="utf-8", errors="replace")
    return _assemble(p.name, [{"text": text, "page": None, "from_ocr": False}])


def _load_docx(p: Path) -> Document:
    import docx  # python-docx

    document = docx.Document(str(p))
    # One block per non-empty paragraph keeps natural structure for segmentation.
    raw = [{"text": para.text, "page": None, "from_ocr": False} for para in document.paragraphs]
    return _assemble(p.name, raw)


def _load_pdf(p: Path) -> Document:
    import pdfplumber

    cfg = get_config()["ingestion"]
    min_chars = cfg["min_chars_per_page_for_text"]
    ocr_enabled = cfg["ocr_enabled"]
    ocr_lang = cfg["ocr_language"]

    raw: list[dict] = []
    with pdfplumber.open(str(p)) as pdf:
        for page_no, page in enumerate(pdf.pages, start=1):
            text = (page.extract_text() or "").strip()
            from_ocr = False

            # Sparse text on a page => likely scanned/image-only => try OCR.
            if len(text) < min_chars and ocr_enabled:
                ocr_text = _ocr_pdf_page(p, page_no, ocr_lang)
                if ocr_text:
                    text, from_ocr = ocr_text, True
                elif not tesseract_available():
                    logger.warning(
                        "Page %d of %s looks scanned but Tesseract is not installed; "
                        "page text will be empty.",
                        page_no,
                        p.name,
                    )

            raw.append({"text": text, "page": page_no, "from_ocr": from_ocr})

    return _assemble(p.name, raw)


def _ocr_pdf_page(p: Path, page_no: int, lang: str) -> str:
    """Render one PDF page to an image and OCR it. Returns "" on any failure
    (missing Poppler/Tesseract, render error) so ingestion degrades gracefully."""
    try:
        from pdf2image import convert_from_path

        images = convert_from_path(str(p), dpi=300, first_page=page_no, last_page=page_no)
    except Exception as exc:  # Poppler missing, corrupt page, etc.
        logger.warning("Could not rasterize page %d of %s for OCR: %s", page_no, p.name, exc)
        return ""

    if not images:
        return ""
    return ocr_image(images[0], lang=lang)
